import io
import re
import json
from fastapi import APIRouter, HTTPException, UploadFile, File, Depends
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from typing import List
from app.core.database import get_db
from app.core.deps import TeacherUser, CurrentUser
from app.models.test import Test
from app.models.question import Question
from app.models.result import Result
from app.models.user import User

router = APIRouter(prefix="/import-export", tags=["import_export"])


# --- Text block extraction -------------------------------------------------
# A "block" is a tuple (text, marked) where `marked` is True when the source
# formatting (e.g. a bold/underlined/coloured run in DOCX) marks the line as
# the correct answer. Non-formatted sources always yield marked=False.

Block = tuple  # (str, bool)

_QUESTION_MARKERS = ("Q:", "П:", "?:")
_ANSWER_MARKERS = ("A:", "В:", "О:")
_CORRECT_CHARS = ("*", "+", "=", "✓", "✔", "√")

# end-of-line cues that identify a question even without explicit markers
_QUESTION_TAIL_RE = re.compile(
    r"(\?|:)\s*$|(означає|розшифров\w*|—\s*це|–\s*це|-\s*це)\s*[:\.]?\s*$",
    re.IGNORECASE,
)
_NUMBER_PREFIX_RE = re.compile(r"^\s*\d{1,3}\s*[\.\)]\s+")
_OPTION_PREFIX_RE = re.compile(r"^\s*([a-zA-Zа-яА-ЯіІїЇєЄ]|\d{1,2})\s*[\.\)]\s+")
_BULLET_PREFIX_RE = re.compile(r"^\s*[-•·–—*]\s+")


def _good_char(c: str) -> bool:
    o = ord(c)
    if c.isspace():
        return True
    if 0x30 <= o <= 0x39:
        return True
    if 0x41 <= o <= 0x5A or 0x61 <= o <= 0x7A:
        return True
    if 0x0400 <= o <= 0x04FF:  # Cyrillic
        return True
    return c in ".,;:?!()[]{}/\\-—–’'\"`+=*%№#@&|<>°…÷×’“”"


def _is_natural_text(s: str) -> bool:
    """True if the string looks like real Latin/Cyrillic prose rather than
    binary garbage that leaked out of a legacy .doc formatting table."""
    if not s or len(s) > 600:
        return False
    good = sum(1 for c in s if _good_char(c))
    if good / len(s) < 0.8:
        return False
    return any(0x41 <= ord(c) <= 0x7A or 0x0400 <= ord(c) <= 0x04FF for c in s)


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1251"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def _extract_txt_blocks(content: bytes) -> List[Block]:
    return [(ln.strip(), False) for ln in _decode_text(content).splitlines() if ln.strip()]


def _extract_pdf_blocks(content: bytes) -> List[Block]:
    import pdfplumber
    blocks: List[Block] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            for ln in text.splitlines():
                ln = ln.strip()
                if ln:
                    blocks.append((ln, False))
    return blocks


def _run_is_marked(run) -> bool:
    try:
        if run.bold or run.underline:
            return True
        color = getattr(run.font, "color", None)
        if color is not None and color.rgb is not None:
            rgb = str(color.rgb)
            if rgb not in ("000000", "auto"):
                return True
    except Exception:
        pass
    return False


def _extract_docx_blocks(content: bytes) -> List[Block]:
    from docx import Document
    doc = Document(io.BytesIO(content))
    blocks: List[Block] = []

    def emit(paragraph_text: str, runs) -> None:
        text = (paragraph_text or "").strip()
        if not text:
            return
        marked = False
        sized = [r for r in runs if r.text and r.text.strip()]
        if sized:
            total = sum(len(r.text) for r in sized)
            mlen = sum(len(r.text) for r in sized if _run_is_marked(r))
            if total and mlen / total > 0.6:
                marked = True
        blocks.append((text, marked))

    for p in doc.paragraphs:
        emit(p.text, p.runs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    emit(p.text, p.runs)
    return blocks


def _extract_doc_blocks(content: bytes) -> List[Block]:
    """Extract paragraphs from a legacy (OLE2) .doc WordDocument stream.

    The stream stores body text either as UTF-16LE or as an 8-bit codepage.
    We try both and keep whichever yields more natural-looking paragraphs.
    Formatting (bold/colour) is not recoverable here, so marked is always
    False and correct answers fall back to needs-review."""
    import olefile
    if not olefile.isOleFile(io.BytesIO(content)):
        return []
    ole = olefile.OleFileIO(io.BytesIO(content))
    if not ole.exists("WordDocument"):
        return []
    data = ole.openstream("WordDocument").read()

    def paragraphs_from(text: str) -> List[str]:
        out: List[str] = []
        for chunk in re.split(r"[\r\x07\x0b\x0c]", text):
            cleaned = "".join(c for c in chunk if c.isprintable() or c == " ").strip()
            if len(cleaned) > 1 and _is_natural_text(cleaned):
                out.append(cleaned)
        return out

    def score(paras: List[str]):
        # A wrong codepage applied to a UTF-16LE stream maps Cyrillic low-bytes
        # onto ASCII letters/punctuation, so ASCII counts are unreliable.
        # Weigh Cyrillic heavily and use letter density as a tiebreaker
        # (real prose is dense; garbled "C o p y" is ~50% spaces).
        text = "".join(paras)
        if not text:
            return (0, 0.0)
        cyr = sum(1 for c in text if 0x0400 <= ord(c) <= 0x04FF)
        letters = sum(1 for c in text if c.isalpha())
        density = letters / len(text)
        return (cyr, density)

    candidates = []
    for enc in ("utf-16-le", "cp1251", "utf-8"):
        try:
            candidates.append(paragraphs_from(data.decode(enc, errors="ignore")))
        except Exception:
            candidates.append([])
    best = max(candidates, key=score)
    return [(p, False) for p in best]


# --- Question parsing -------------------------------------------------------

def _is_question_line(text: str) -> bool:
    if _QUESTION_TAIL_RE.search(text):
        return True
    if _NUMBER_PREFIX_RE.match(text) and len(text) > 12:
        return True
    return False


def _clean_question(text: str) -> str:
    for m in _QUESTION_MARKERS:
        if text.startswith(m):
            text = text[len(m):]
            break
    text = _NUMBER_PREFIX_RE.sub("", text)
    return text.strip()


def _clean_option(text: str):
    """Return (clean_text, is_correct) stripping option markers/enumerators."""
    is_correct = False
    t = text.strip()
    for m in _ANSWER_MARKERS:
        if t.startswith(m):
            t = t[len(m):].strip()
            break
    while t and t[0] in _CORRECT_CHARS:
        is_correct = True
        t = t[1:].strip()
    if t.startswith("(") and t[1:2] in ("+", "✓"):
        is_correct = True
        t = t[t.find(")") + 1:].strip() if ")" in t else t
    t = _OPTION_PREFIX_RE.sub("", t)
    t = _BULLET_PREFIX_RE.sub("", t)
    return t.strip(), is_correct


def _finalize_group(content: str, options: List[tuple]) -> dict | None:
    """options: list of (text, is_correct). Build a question dict or None."""
    choices = [o[0] for o in options if o[0]]
    if len(choices) < 2:
        return None
    marked = [o[0] for o in options if o[1] and o[0]]
    needs_review = False
    if len(marked) == 1:
        correct = marked[0]
        qtype = "single-choice"
    elif 2 <= len(marked) < len(choices):
        correct = marked
        qtype = "multiple-choice"
    else:
        # zero marks, or everything marked (unreliable) -> default + review
        correct = choices[0]
        qtype = "single-choice"
        needs_review = True
    return {
        "type": qtype,
        "content": content,
        "options": {"choices": choices, "correct": correct},
        "needs_review": needs_review,
    }


def _parse_blocks(blocks: List[Block]) -> List[dict]:
    texts = [b[0] for b in blocks]
    has_markers = any(
        t.startswith(_QUESTION_MARKERS) or t.startswith(_ANSWER_MARKERS) for t in texts
    )

    questions: List[dict] = []
    content: str | None = None
    options: List[tuple] = []

    def flush():
        nonlocal content, options
        if content is not None:
            q = _finalize_group(content, options)
            if q:
                questions.append(q)
        content, options = None, []

    if has_markers:
        for text, marked in blocks:
            if text.startswith(_QUESTION_MARKERS):
                flush()
                content = _clean_question(text)
                options = []
            elif content is not None:
                clean, is_corr = _clean_option(text)
                if clean:
                    options.append((clean, is_corr or marked))
        flush()
        return questions

    # markerless / numbered format -------------------------------------------
    # Pass 1: group by question cues (lines ending in ? : or "— це"/"означає").
    groups: List[dict] = []
    cur: dict | None = None
    for text, marked in blocks:
        if _is_question_line(text):
            if cur and len([o for o in cur["opts"] if o[0]]) >= 2:
                groups.append(cur)
            cur = {"content": _clean_question(text), "opts": []}
        elif cur is not None:
            clean, is_corr = _clean_option(text)
            if clean:
                cur["opts"].append((clean, is_corr or marked))
    if cur and len([o for o in cur["opts"] if o[0]]) >= 2:
        groups.append(cur)

    # Determine the modal options-per-question. Many real test banks use a
    # fixed layout (e.g. exactly 4 options) but only some question lines end
    # with a "?", so statement-style questions get absorbed as extra options.
    from collections import Counter
    counts = Counter(len(g["opts"]) for g in groups)
    modal = counts.most_common(1)[0][0] if counts else 0

    # Pass 2: re-split oversized groups that clearly contain >=2 questions.
    for g in groups:
        opts = g["opts"]
        if modal >= 3 and len(opts) >= 2 * modal:
            seq = [(g["content"], False)] + opts
            i = 0
            while i < len(seq):
                qc = seq[i][0]
                block = seq[i + 1: i + 1 + modal]
                if len([b for b in block if b[0]]) >= 2:
                    q = _finalize_group(qc, block)
                    if q:
                        questions.append(q)
                i += 1 + modal
        else:
            q = _finalize_group(g["content"], opts)
            if q:
                questions.append(q)
    return questions


def _parse_file(filename: str, content: bytes) -> List[dict]:
    """Dispatch parsing based on file extension. Raises HTTPException for
    unsupported formats."""
    name = (filename or "").lower()
    if name.endswith(".docx"):
        try:
            return _parse_blocks(_extract_docx_blocks(content))
        except Exception:
            return []
    if name.endswith(".doc"):
        # some .doc files are actually .docx in disguise
        try:
            return _parse_blocks(_extract_docx_blocks(content))
        except Exception:
            pass
        try:
            return _parse_blocks(_extract_doc_blocks(content))
        except Exception:
            return []
    if name.endswith(".xlsx"):
        return _parse_xlsx(content)
    if name.endswith(".txt"):
        return _parse_blocks(_extract_txt_blocks(content))
    if name.endswith(".pdf"):
        try:
            return _parse_blocks(_extract_pdf_blocks(content))
        except Exception:
            return []
    if name.endswith(".json"):
        data = json.loads(content)
        if isinstance(data, dict) and "questions" in data:
            data = data["questions"]
        return data if isinstance(data, list) else []
    raise HTTPException(status_code=400, detail="Unsupported file format")


def _normalize_questions(raw_questions: List[dict]) -> List[dict]:
    normalized: List[dict] = []
    for q in raw_questions:
        if not isinstance(q, dict):
            continue
        normalized.append({
            "type": q.get("type", "single-choice"),
            "content": q.get("content", ""),
            "hint": q.get("hint"),
            "explanation": q.get("explanation"),
            "points": q.get("points", 1),
            "options": q.get("options", {}),
            "needs_review": bool(q.get("needs_review", False)),
        })
    return normalized


def _parse_xlsx(content: bytes) -> List[dict]:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(content))
        ws = wb.active
        questions = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not row[0]:
                continue
            q_type = str(row[1]).strip() if row[1] else "single-choice"
            choices_raw = str(row[3]).strip() if row[3] else ""
            choices = [c.strip() for c in choices_raw.split("|") if c.strip()]
            correct_raw = str(row[4]).strip() if row[4] else ""
            questions.append({
                "type": q_type,
                "content": str(row[0]).strip(),
                "hint": str(row[2]).strip() if row[2] else None,
                "options": {"choices": choices, "correct": correct_raw},
            })
        return questions
    except Exception:
        return []


@router.post("/tests/{test_id}/import")
async def import_questions(
    test_id: str,
    current_user: TeacherUser,
    db: AsyncSession = Depends(get_db),
    file: UploadFile = File(...),
):
    test_r = await db.execute(select(Test).where(Test.id == test_id))
    test = test_r.scalar_one_or_none()
    if not test:
        raise HTTPException(status_code=404, detail="Test not found")

    content = await file.read()
    try:
        questions_data = _parse_file(file.filename or "", content)
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=400, detail="Не вдалося прочитати файл")

    max_idx_r = await db.execute(
        select(Question).where(Question.test_id == test_id).order_by(Question.order_index.desc()).limit(1)
    )
    last_q = max_idx_r.scalar_one_or_none()
    start_idx = (last_q.order_index + 1) if last_q else 0

    added = 0
    for i, qdata in enumerate(questions_data):
        q = Question(
            test_id=test_id,
            type=qdata.get("type", "single-choice"),
            content=qdata.get("content", ""),
            hint=qdata.get("hint"),
            explanation=qdata.get("explanation"),
            points=qdata.get("points", 1),
            order_index=start_idx + i,
            options=qdata.get("options", {}),
        )
        db.add(q)
        added += 1

    return {"imported": added}


@router.post("/parse")
async def parse_file(
    current_user: CurrentUser,
    file: UploadFile = File(...),
):
    """Parse an uploaded file (JSON, DOCX, DOC, XLSX, TXT, PDF) and return the
    extracted questions WITHOUT persisting anything. Used by the test-creation
    UI to preview/edit questions before saving the test."""
    content = await file.read()
    try:
        questions_data = _parse_file(file.filename or "", content)
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=400, detail="Не вдалося прочитати файл")

    normalized = _normalize_questions(questions_data)
    return {"count": len(normalized), "questions": normalized}


@router.get("/tests/{test_id}/export/pdf")
async def export_test_pdf(
    test_id: str,
    current_user: TeacherUser,
    db: AsyncSession = Depends(get_db),
):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    test_r = await db.execute(select(Test).where(Test.id == test_id))
    test = test_r.scalar_one_or_none()
    if not test:
        raise HTTPException(status_code=404, detail="Test not found")

    q_r = await db.execute(
        select(Question).where(Question.test_id == test_id).order_by(Question.order_index)
    )
    questions = q_r.scalars().all()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    styles = getSampleStyleSheet()
    story = [Paragraph(test.title, styles["Title"]), Spacer(1, 12)]

    for i, q in enumerate(questions, 1):
        story.append(Paragraph(f"{i}. {q.content}", styles["Normal"]))
        choices = q.options.get("choices", [])
        for ch in choices:
            story.append(Paragraph(f"    ○ {ch}", styles["Normal"]))
        story.append(Spacer(1, 6))

    doc.build(story)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={test.title}.pdf"},
    )


@router.get("/results/{result_id}/certificate.pdf")
async def student_certificate(
    result_id: str,
    current_user: CurrentUser,
    db: AsyncSession = Depends(get_db),
):
    """Personal PDF certificate/report for a completed test attempt. Students
    can only download their own; teachers/admins can fetch any."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    result_r = await db.execute(select(Result).where(Result.id == result_id))
    result = result_r.scalar_one_or_none()
    if not result:
        raise HTTPException(status_code=404, detail="Result not found")
    if result.user_id != current_user.id and current_user.role not in ("teacher", "admin"):
        raise HTTPException(status_code=403, detail="Not authorized")

    user_r = await db.execute(select(User).where(User.id == result.user_id))
    user = user_r.scalar_one_or_none()
    test_r = await db.execute(select(Test).where(Test.id == result.test_id))
    test = test_r.scalar_one_or_none()
    if not user or not test:
        raise HTTPException(status_code=404, detail="Associated user or test missing")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "title", parent=styles["Title"], fontSize=24, leading=30,
        textColor=HexColor("#4c1d95"), alignment=TA_CENTER,
    )
    subtitle_style = ParagraphStyle(
        "subtitle", parent=styles["Normal"], fontSize=12,
        textColor=HexColor("#6b7280"), alignment=TA_CENTER, leading=18,
    )
    huge_style = ParagraphStyle(
        "huge", parent=styles["Normal"], fontSize=48, leading=56,
        textColor=HexColor("#059669") if result.passed else HexColor("#e11d48"),
        alignment=TA_CENTER,
    )
    label_style = ParagraphStyle(
        "label", parent=styles["Normal"], fontSize=10,
        textColor=HexColor("#9ca3af"), alignment=TA_CENTER,
    )
    name_style = ParagraphStyle(
        "name", parent=styles["Normal"], fontSize=18,
        textColor=HexColor("#111827"), alignment=TA_CENTER, leading=24,
    )

    story = []
    story.append(Paragraph("Сертифікат проходження тесту", title_style))
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph("TestAP — Платформа інтелектуального тестування", subtitle_style))
    story.append(Spacer(1, 1.4 * cm))

    story.append(Paragraph("Цей сертифікат підтверджує, що", label_style))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(user.full_name, name_style))
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph("успішно пройшов(ла) тест" if result.passed else "пройшов(ла) тест", subtitle_style))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(f"<b>{test.title}</b>", name_style))
    story.append(Spacer(1, 1.2 * cm))

    story.append(Paragraph(f"{round(result.percent)}%", huge_style))
    story.append(Spacer(1, 0.2 * cm))
    status_text = "ТЕСТ ПРОЙДЕНО" if result.passed else "ТЕСТ НЕ ПРОЙДЕНО"
    story.append(Paragraph(status_text, ParagraphStyle(
        "status", parent=styles["Normal"], fontSize=14, alignment=TA_CENTER,
        textColor=HexColor("#059669") if result.passed else HexColor("#e11d48"),
    )))
    story.append(Spacer(1, 1.2 * cm))

    rows = [
        ["Бали", f"{result.score} / {result.max_score}"],
        ["Прохідний бал", f"{test.passing_score}%"],
        ["Дата проходження", result.completed_at.strftime("%d.%m.%Y %H:%M")],
        ["Email учасника", user.email],
        ["ID сертифіката", result.id],
    ]
    table = Table(rows, colWidths=[6 * cm, 9 * cm])
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("TEXTCOLOR", (0, 0), (0, -1), HexColor("#6b7280")),
        ("TEXTCOLOR", (1, 0), (1, -1), HexColor("#111827")),
        ("ALIGN", (0, 0), (0, -1), "RIGHT"),
        ("ALIGN", (1, 0), (1, -1), "LEFT"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("LINEBELOW", (0, 0), (-1, -1), 0.4, HexColor("#e5e7eb")),
    ]))
    story.append(table)
    story.append(Spacer(1, 1.5 * cm))
    story.append(Paragraph("© TestAP — Розумне тестування для освіти", label_style))

    doc.build(story)
    buf.seek(0)
    filename = f"certificate-{result.id[:8]}.pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f"inline; filename={filename}"},
    )


@router.get("/tests/{test_id}/results/export/xlsx")
async def export_results_xlsx(
    test_id: str,
    current_user: TeacherUser,
    db: AsyncSession = Depends(get_db),
):
    import openpyxl
    results_r = await db.execute(
        select(Result, User.full_name, User.email)
        .join(User, User.id == Result.user_id)
        .where(Result.test_id == test_id)
        .order_by(Result.completed_at.desc())
    )
    rows = results_r.all()

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Results"
    ws.append(["Ім'я", "Email", "Бал", "Максимум", "%", "Статус", "Дата"])
    for res, name, email in rows:
        ws.append([
            name, email, res.score, res.max_score, res.percent,
            "Пройшов" if res.passed else "Не пройшов",
            res.completed_at.strftime("%Y-%m-%d %H:%M"),
        ])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=results.xlsx"},
    )
